#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
SQL Server 数据库健康检查脚本
功能：连接 SQL Server 数据库，查询各项指标，评估健康状态，生成 Word 报告
"""

import pyodbc
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime
import os
import re


class SQLServerHealthChecker:
    """SQL Server 数据库健康检查器"""
    
    def __init__(self, server, database, username, password):
        """
        初始化连接参数
        :param server: 服务器地址及端口
        :param database: 数据库名
        :param username: 用户名
        :param password: 密码
        """
        self.server = server
        self.database = database
        self.username = username
        self.password = password
        self.connection = None
        self.health_issues = []
        self.metrics = {}
        
    def connect(self):
        """建立数据库连接"""
        try:
            conn_str = (
                f"DRIVER={{ODBC Driver 17 for SQL Server}};"
                f"SERVER={self.server};"
                f"DATABASE={self.database};"
                f"UID={self.username};"
                f"PWD={self.password}"
            )
            self.connection = pyodbc.connect(conn_str)
            print(f"✓ 成功连接到数据库：{self.server}/{self.database}")
            return True
        except Exception as e:
            print(f"✗ 连接失败：{str(e)}")
            return False
    
    def disconnect(self):
        """断开数据库连接"""
        if self.connection:
            self.connection.close()
            print("✓ 已断开数据库连接")
    
    def execute_query(self, query):
        """执行查询并返回结果"""
        try:
            cursor = self.connection.cursor()
            cursor.execute(query)
            columns = [column[0] for column in cursor.description]
            results = []
            for row in cursor.fetchall():
                results.append(dict(zip(columns, row)))
            return results
        except Exception as e:
            print(f"查询执行失败：{str(e)}")
            return []
    
    def check_database_info(self):
        """查询数据库基本信息"""
        query = """
        SELECT 
            DB_NAME() AS database_name,
            @@VERSION AS sql_server_version,
            SERVERPROPERTY('ProductLevel') AS product_level,
            SERVERPROPERTY('Edition') AS edition,
            SERVERPROPERTY('MachineName') AS machine_name,
            SERVERPROPERTY('InstanceName') AS instance_name,
            SERVERPROPERTY('IsClustered') AS is_clustered
        """
        results = self.execute_query(query)
        if results:
            self.metrics['database_info'] = results[0]
        return results
    
    def check_database_size(self):
        """查询所有数据库大小 (使用 sys.master_files)"""
        query = """
        SELECT 
            DB_NAME(database_id) AS database_name,
            SUM(size * 8.0 / 1024) AS total_size_mb,
            SUM(CASE WHEN type = 0 THEN size * 8.0 / 1024 ELSE 0 END) AS data_size_mb,
            SUM(CASE WHEN type = 1 THEN size * 8.0 / 1024 ELSE 0 END) AS log_size_mb
        FROM sys.master_files
        GROUP BY database_id, DB_NAME(database_id)
        ORDER BY total_size_mb DESC
        """
        results = self.execute_query(query)
        self.metrics['database_size'] = results  # 现在存储的是所有数据库的列表
        
        # 评估所有数据库的大小
        for row in results:
            size_mb = row.get('total_size_mb', 0)
            db_name = row.get('database_name', '')
            if size_mb > 100000:
                self.health_issues.append({
                    'level': '警告',
                    'item': '数据库大小',
                    'description': f'数据库 [{db_name}] 大小超过 100GB (当前：{size_mb:.2f} MB)',
                    'suggestion': '考虑归档历史数据或优化存储'
                })
        return results
    
    def check_table_space(self):
        """查询表空间使用情况"""
        query = """
        SELECT TOP 20
            t.name AS table_name,
            SUM(a.total_pages) * 8 AS total_space_kb,
            SUM(a.used_pages) * 8 AS used_space_kb,
            (SUM(a.total_pages) - SUM(a.used_pages)) * 8 AS unused_space_kb
        FROM sys.tables t
        INNER JOIN sys.indexes i ON t.OBJECT_ID = i.object_id
        INNER JOIN sys.partitions p ON i.object_id = p.OBJECT_ID AND i.index_id = p.index_id
        INNER JOIN sys.allocation_units a ON p.partition_id = a.container_id
        GROUP BY t.name
        ORDER BY total_space_kb DESC
        """
        results = self.execute_query(query)
        self.metrics['table_space'] = results
        return results
    
    def check_index_fragmentation(self):
        """查询索引碎片情况"""
        query = """
        SELECT TOP 20
            OBJECT_NAME(ips.object_id) AS table_name,
            i.name AS index_name,
            ips.avg_fragmentation_in_percent,
            ips.page_count
        FROM sys.dm_db_index_physical_stats(DB_ID(), NULL, NULL, NULL, 'LIMITED') ips
        INNER JOIN sys.indexes i ON ips.object_id = i.object_id AND ips.index_id = i.index_id
        WHERE ips.avg_fragmentation_in_percent > 10 AND ips.page_count > 100
        ORDER BY ips.avg_fragmentation_in_percent DESC
        """
        results = self.execute_query(query)
        self.metrics['index_fragmentation'] = results
        
        for row in results:
            frag = row.get('avg_fragmentation_in_percent', 0)
            if frag > 30:
                self.health_issues.append({
                    'level': '严重',
                    'item': '索引碎片',
                    'description': f"表 {row.get('table_name')} 的索引 {row.get('index_name')} 碎片率：{frag:.2f}%",
                    'suggestion': '建议执行索引重组或重建'
                })
            elif frag > 10:
                self.health_issues.append({
                    'level': '警告',
                    'item': '索引碎片',
                    'description': f"表 {row.get('table_name')} 的索引 {row.get('index_name')} 碎片率：{frag:.2f}%",
                    'suggestion': '考虑执行索引重组'
                })
        return results
    
    def check_missing_indexes(self):
        """查询缺失的索引"""
        query = """
        SELECT TOP 10
            OBJECT_NAME(mid.object_id) AS table_name,
            mid.equality_columns,
            mid.inequality_columns,
            mid.included_columns,
            migs.avg_user_impact AS impact,
            migs.user_seeks,
            migs.user_scans
        FROM sys.dm_db_missing_index_details mid
        INNER JOIN sys.dm_db_missing_index_groups mig ON mid.index_handle = mig.index_handle
        INNER JOIN sys.dm_db_missing_index_group_stats migs ON mig.index_group_handle = migs.group_handle
        WHERE migs.avg_user_impact > 30
        ORDER BY migs.avg_user_impact DESC
        """
        results = self.execute_query(query)
        self.metrics['missing_indexes'] = results
        
        for row in results:
            impact = row.get('impact', 0)
            if impact > 70:
                self.health_issues.append({
                    'level': '严重',
                    'item': '缺失索引',
                    'description': f"表 {row.get('table_name')} 缺失索引，影响度：{impact}%",
                    'suggestion': '建议创建索引以提升查询性能'
                })
            elif impact > 30:
                self.health_issues.append({
                    'level': '警告',
                    'item': '缺失索引',
                    'description': f"表 {row.get('table_name')} 缺失索引，影响度：{impact}%",
                    'suggestion': '考虑创建索引'
                })
        return results
    
    def check_wait_statistics(self):
        """查询等待统计"""
        query = """
        SELECT TOP 10
            wait_type,
            waiting_tasks_count,
            wait_time_ms / 1000.0 AS wait_time_seconds,
            max_wait_time_ms / 1000.0 AS max_wait_seconds
        FROM sys.dm_os_wait_stats
        WHERE wait_type NOT IN (
            'BROKER_EVENTACTIVATOR', 'BROKER_RECEIVE_WAITFOR', 'BROKER_TASK_STOP',
            'CLR_AUTO_EVENT', 'CLR_MANUAL_EVENT', 'CLR_SEMAPHORE_2',
            'DBMIRROR_DBM_EVENT', 'DBMIRROR_EVENTS_QUEUE',
            'DISPATCHER_QUEUE_SEMAPHORE', 'FT_IFTS_SCHEDULER_IDLE_WAIT',
            'LOGMGR_QUEUE', 'REDO_THREAD_PENDING_WORK',
            'REQUEST_FOR_DEADLOCK_SEARCH', 'SERVER_IDLE_CHECK',
            'SLEEP_TASK', 'SQLTRACE_BUFFER_FLUSH', 'SQLTRACE_INCREMENTAL_FLUSH_SLEEP',
            'WAITFOR', 'XE_DISPATCHER_JOIN', 'XE_DISPATCHER_WAIT',
            'XE_TIMER_EVENT', 'BROKER_TO_FLUSH', 'HADR_FILESTREAM_I_COMPLETE',
            'HADR_CLUSAPI_CALL', 'HADR_LOGCAPTURE_WAIT', 'HADR_NOTIFICATION_DEQUEUE',
            'HADR_WORK_QUEUE', 'KSOURCE_WAKEUP', 'LAZYWRITER_SLEEP',
            'ONDEMAND_TASK_QUEUE', 'PWAIT_ALL_COMPONENTS_INITIALIZED',
            'PWAIT_DIRECTLOGCONSUMER_GETNEXT', 'QDS_PERSIST_TASK_MAIN_LOOP_SLEEP',
            'QDS_ASYNC_QUEUE', 'QDS_CLEANUP_STALE_QUERIES_TASK_MAIN_LOOP_SLEEP',
            'QDS_SHUTDOWN_QUEUE', 'REDO_THREAD_PENDING_WORK',
            'RESOURCE_QUEUE', 'SOFS_LOCK_DATABASE', 'SOS_WORK_DISPATCHER',
            'SPVD_SERVER_CACHECLEANUP', 'UCS_SESSION_REGISTRATION',
            'WAIT_XTP_OFFLINE_CKPT_NEW_LOG'
        )
        ORDER BY wait_time_ms DESC
        """
        results = self.execute_query(query)
        self.metrics['wait_statistics'] = results
        
        critical_waits = ['PAGEIOLATCH', 'WRITELOG', 'LCK_M', 'CXPACKET', 'SOS_SCHEDULER_YIELD']
        for row in results:
            wait_type = row.get('wait_type', '')
            wait_time = row.get('wait_time_seconds', 0)
            for critical in critical_waits:
                if critical in wait_type and wait_time > 60:
                    self.health_issues.append({
                        'level': '警告',
                        'item': '等待统计',
                        'description': f"等待类型 {wait_type} 累计等待时间：{wait_time:.2f}秒",
                        'suggestion': '分析性能瓶颈，优化相关查询或硬件配置'
                    })
        return results
    
    def check_active_sessions(self):
        """查询活跃会话"""
        query = """
        SELECT 
            COUNT(*) AS total_sessions,
            SUM(CASE WHEN status = 'running' THEN 1 ELSE 0 END) AS running_sessions,
            SUM(CASE WHEN status = 'sleeping' THEN 1 ELSE 0 END) AS sleeping_sessions,
            SUM(CASE WHEN status = 'suspended' THEN 1 ELSE 0 END) AS suspended_sessions
        FROM sys.dm_exec_sessions
        WHERE is_user_process = 1
        """
        results = self.execute_query(query)
        if results:
            self.metrics['active_sessions'] = results[0]
        return results
    
    def check_cpu_memory(self):
        """查询 CPU 和内存使用情况"""
        query = """
        SELECT 
            (SELECT cntr_value FROM sys.dm_os_performance_counters 
            WHERE object_name = 'SQLServer:Memory Manager' AND counter_name = 'Total Server Memory (KB)') AS total_memory_kb,
            (SELECT cntr_value FROM sys.dm_os_performance_counters 
            WHERE object_name = 'SQLServer:Memory Manager' AND counter_name = 'Target Server Memory (KB)') AS target_memory_kb,
            (SELECT cntr_value FROM sys.dm_os_performance_counters 
            WHERE object_name = 'SQLServer:Buffer Manager' AND counter_name = 'Page life expectancy') AS page_life_expectancy;
        """
        results = self.execute_query(query)
        if results:
            self.metrics['cpu_memory'] = results[0]
            ple = results[0].get('page_life_expectancy', 0)
            if ple < 300:
                self.health_issues.append({
                    'level': '严重',
                    'item': '内存性能',
                    'description': f'页面预期寿命 (PLE) 过低：{ple}秒',
                    'suggestion': '增加服务器内存或优化查询以减少内存压力'
                })
            elif ple < 600:
                self.health_issues.append({
                    'level': '警告',
                    'item': '内存性能',
                    'description': f'页面预期寿命 (PLE) 偏低：{ple}秒',
                    'suggestion': '关注内存使用情况，考虑优化'
                })
        return results
    
    def check_backup_status(self):
        """查询备份状态"""
        query = """
        SELECT TOP 10
            bs.database_name,
            bs.type AS backup_type,
            bs.backup_start_date,
            bs.backup_finish_date,
            DATEDIFF(SECOND, bs.backup_start_date, bs.backup_finish_date) AS duration_seconds,
            bs.backup_size / 1024 / 1024 AS backup_size_mb,
            bmf.physical_device_name
        FROM msdb.dbo.backupset bs
        INNER JOIN msdb.dbo.backupmediafamily bmf ON bs.media_set_id = bmf.media_set_id
        WHERE bs.database_name = DB_NAME()
        ORDER BY bs.backup_start_date DESC
        """
        results = self.execute_query(query)
        self.metrics['backup_status'] = results
        
        if results:
            last_backup = results[0].get('backup_start_date')
            if last_backup:
                hours_since_backup = (datetime.now() - last_backup).total_seconds() / 3600
                if hours_since_backup > 24:
                    self.health_issues.append({
                        'level': '严重',
                        'item': '备份状态',
                        'description': f'距离上次备份已超过 {hours_since_backup:.1f} 小时',
                        'suggestion': '立即执行数据库备份'
                    })
        else:
            self.health_issues.append({
                'level': '严重',
                'item': '备份状态',
                'description': '未找到备份记录',
                'suggestion': '立即配置并执行数据库备份'
            })
        return results
    
    def check_database_files(self):
        """查询数据库文件状态"""
        query = """
        SELECT 
            name AS file_name,
            type_desc AS file_type,
            size * 8.0 / 1024 AS size_mb,
            max_size * 8.0 / 1024 AS max_size_mb,
            growth,
            is_percent_growth,
            state_desc
        FROM sys.database_files
        """
        results = self.execute_query(query)
        self.metrics['database_files'] = results
        
        for row in results:
            state = row.get('state_desc', '')
            if state != 'ONLINE':
                self.health_issues.append({
                    'level': '严重',
                    'item': '文件状态',
                    'description': f"文件 {row.get('file_name')} 状态异常：{state}",
                    'suggestion': '立即检查文件状态并修复'
                })
        return results
    
    def check_locks(self):
        """查询锁等待情况"""
        query = """
        SELECT 
            COUNT(*) AS lock_count,
            DB_NAME(resource_database_id) AS database_name
        FROM sys.dm_tran_locks
        WHERE request_status = 'WAIT'
        GROUP BY DB_NAME(resource_database_id)
        """
        results = self.execute_query(query)
        self.metrics['locks'] = results
        
        for row in results:
            lock_count = row.get('lock_count', 0)
            if lock_count > 10:
                self.health_issues.append({
                    'level': '警告',
                    'item': '锁等待',
                    'description': f"数据库 {row.get('database_name')} 有 {lock_count} 个锁等待",
                    'suggestion': '检查长时间运行的事务和阻塞链'
                })
        return results
    
    def check_deadlocks(self):
        """查询死锁统计"""
        query = """
        SELECT 
            cntr_value AS deadlock_count
        FROM sys.dm_os_performance_counters
        WHERE counter_name = 'Number of Deadlocks/sec'
        """
        results = self.execute_query(query)
        if results:
            self.metrics['deadlocks'] = results[0]
            deadlock_count = results[0].get('deadlock_count', 0)
            if deadlock_count > 0:
                self.health_issues.append({
                    'level': '警告',
                    'item': '死锁统计',
                    'description': f'检测到 {deadlock_count} 次死锁/秒',
                    'suggestion': '分析死锁图，优化事务和锁策略'
                })
        return results
    
    def run_all_checks(self):
        """执行所有检查"""
        print("\n开始执行数据库健康检查...")
        print("=" * 60)
        
        checks = [
            ('数据库基本信息', self.check_database_info),
            ('所有数据库大小', self.check_database_size),
            ('表空间使用', self.check_table_space),
            ('索引碎片', self.check_index_fragmentation),
            ('缺失索引', self.check_missing_indexes),
            ('等待统计', self.check_wait_statistics),
            ('活跃会话', self.check_active_sessions),
            ('CPU 和内存', self.check_cpu_memory),
            ('备份状态', self.check_backup_status),
            ('数据库文件', self.check_database_files),
            ('锁等待', self.check_locks),
            ('死锁统计', self.check_deadlocks),
        ]
        
        for check_name, check_func in checks:
            try:
                print(f"正在检查：{check_name}...")
                check_func()
            except Exception as e:
                print(f"✗ 检查 {check_name} 失败：{str(e)}")
        
        print("=" * 60)
        print(f"检查完成。发现 {len(self.health_issues)} 个问题。\n")
    
    def get_health_score(self):
        """计算健康评分"""
        if not self.health_issues:
            return 100
        
        score = 100
        for issue in self.health_issues:
            if issue['level'] == '严重':
                score -= 15
            elif issue['level'] == '警告':
                score -= 5
        
        return max(0, score)
    
    def get_health_level(self):
        """获取健康等级"""
        score = self.get_health_score()
        if score >= 90:
            return '优秀', RGBColor(0, 128, 0)
        elif score >= 70:
            return '良好', RGBColor(128, 128, 0)
        elif score >= 50:
            return '中等', RGBColor(255, 165, 0)
        else:
            return '差', RGBColor(255, 0, 0)
    
    def generate_report(self, output_path=None):
        """生成 Word 报告"""
        if output_path is None:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_path = f"SQL_Server_Health_Report_{timestamp}.docx"
        
        doc = Document()
        
        # 标题
        title = doc.add_heading('SQL Server 数据库健康检查报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 基本信息
        doc.add_heading('1. 报告基本信息', level=1)
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        
        info_data = [
            ('检查时间', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
            ('服务器', self.server),
            ('数据库', self.database),
            ('用户', self.username),
        ]
        
        for i, (label, value) in enumerate(info_data):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = str(value)
            table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        
        doc.add_paragraph()
        
        # 健康评分
        doc.add_heading('2. 健康评估结果', level=1)
        score = self.get_health_score()
        level, color = self.get_health_level()
        
        para = doc.add_paragraph()
        para.add_run(f'健康评分：{score} 分\n').bold = True
        para.add_run(f'健康等级：{level}\n').bold = True
        para.runs[1].font.color.rgb = color
        
        doc.add_paragraph()
        
        # 问题汇总
        doc.add_heading('3. 发现的问题', level=1)
        if self.health_issues:
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '级别'
            hdr_cells[1].text = '检查项'
            hdr_cells[2].text = '问题描述'
            hdr_cells[3].text = '建议'
            
            for cell in hdr_cells:
                cell.paragraphs[0].runs[0].bold = True
            
            for issue in self.health_issues:
                row_cells = table.add_row().cells
                row_cells[0].text = issue['level']
                row_cells[1].text = issue['item']
                row_cells[2].text = issue['description']
                row_cells[3].text = issue['suggestion']
                
                if issue['level'] == '严重':
                    row_cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
                elif issue['level'] == '警告':
                    row_cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 165, 0)
        else:
            doc.add_paragraph('未发现明显问题，数据库运行状态良好。')
        
        doc.add_paragraph()
        
        # 详细指标
        doc.add_heading('4. 详细监控指标', level=1)
        
        # 数据库信息
        if 'database_info' in self.metrics:
            doc.add_heading('4.1 数据库基本信息', level=2)
            info = self.metrics['database_info']
            for key, value in info.items():
                if value:
                    doc.add_paragraph(f'{key}: {value}')
            doc.add_paragraph()
        
        # 数据库大小 (修改为表格展示所有数据库)
        if 'database_size' in self.metrics and self.metrics['database_size']:
            doc.add_heading('4.2 所有数据库大小', level=2)
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '数据库名'
            hdr_cells[1].text = '总大小 (MB)'
            hdr_cells[2].text = '数据文件 (MB)'
            hdr_cells[3].text = '日志文件 (MB)'
            
            for cell in hdr_cells:
                cell.paragraphs[0].runs[0].bold = True
            
            for row in self.metrics['database_size']:
                row_cells = table.add_row().cells
                row_cells[0].text = str(row.get('database_name', ''))
                row_cells[1].text = f"{row.get('total_size_mb', 0):.2f}"
                row_cells[2].text = f"{row.get('data_size_mb', 0):.2f}"
                row_cells[3].text = f"{row.get('log_size_mb', 0):.2f}"
            doc.add_paragraph()
        
        # 表空间
        if 'table_space' in self.metrics and self.metrics['table_space']:
            doc.add_heading('4.3 表空间使用 (TOP 20)', level=2)
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '表名'
            hdr_cells[1].text = '总空间 (KB)'
            hdr_cells[2].text = '已用空间 (KB)'
            
            for cell in hdr_cells:
                cell.paragraphs[0].runs[0].bold = True
            
            for row in self.metrics['table_space']:
                row_cells = table.add_row().cells
                row_cells[0].text = str(row.get('table_name', ''))
                row_cells[1].text = str(row.get('total_space_kb', 0))
                row_cells[2].text = str(row.get('used_space_kb', 0))
            doc.add_paragraph()
        
        # 索引碎片
        if 'index_fragmentation' in self.metrics and self.metrics['index_fragmentation']:
            doc.add_heading('4.4 索引碎片 (碎片率>10%)', level=2)
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '表名'
            hdr_cells[1].text = '索引名'
            hdr_cells[2].text = '碎片率 (%)'
            hdr_cells[3].text = '页数'
            
            for cell in hdr_cells:
                cell.paragraphs[0].runs[0].bold = True
            
            for row in self.metrics['index_fragmentation']:
                row_cells = table.add_row().cells
                row_cells[0].text = str(row.get('table_name', ''))
                row_cells[1].text = str(row.get('index_name', ''))
                row_cells[2].text = f"{row.get('avg_fragmentation_in_percent', 0):.2f}"
                row_cells[3].text = str(row.get('page_count', 0))
            doc.add_paragraph()
        
        # 缺失索引
        if 'missing_indexes' in self.metrics and self.metrics['missing_indexes']:
            doc.add_heading('4.5 缺失的索引 (影响度>30%)', level=2)
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '表名'
            hdr_cells[1].text = '等值列'
            hdr_cells[2].text = '不等值列'
            hdr_cells[3].text = '包含列'
            hdr_cells[4].text = '影响度 (%)'
            
            for cell in hdr_cells:
                cell.paragraphs[0].runs[0].bold = True
            
            for row in self.metrics['missing_indexes']:
                row_cells = table.add_row().cells
                row_cells[0].text = str(row.get('table_name', ''))
                row_cells[1].text = str(row.get('equality_columns', '') or 'N/A')
                row_cells[2].text = str(row.get('inequality_columns', '') or 'N/A')
                row_cells[3].text = str(row.get('included_columns', '') or 'N/A')
                row_cells[4].text = str(row.get('impact', 0))
            doc.add_paragraph()
        
        # 等待统计
        if 'wait_statistics' in self.metrics and self.metrics['wait_statistics']:
            doc.add_heading('4.6 等待统计 (TOP 10)', level=2)
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '等待类型'
            hdr_cells[1].text = '等待次数'
            hdr_cells[2].text = '总等待时间 (秒)'
            hdr_cells[3].text = '最大等待时间 (秒)'
            
            for cell in hdr_cells:
                cell.paragraphs[0].runs[0].bold = True
            
            for row in self.metrics['wait_statistics']:
                row_cells = table.add_row().cells
                row_cells[0].text = str(row.get('wait_type', ''))
                row_cells[1].text = str(row.get('waiting_tasks_count', 0))
                row_cells[2].text = f"{row.get('wait_time_seconds', 0):.2f}"
                row_cells[3].text = f"{row.get('max_wait_seconds', 0):.2f}"
            doc.add_paragraph()
        
        # 活跃会话
        if 'active_sessions' in self.metrics:
            doc.add_heading('4.7 活跃会话', level=2)
            sessions = self.metrics['active_sessions']
            doc.add_paragraph(f"总会话数：{sessions.get('total_sessions', 0)}")
            doc.add_paragraph(f"运行中：{sessions.get('running_sessions', 0)}")
            doc.add_paragraph(f"睡眠中：{sessions.get('sleeping_sessions', 0)}")
            doc.add_paragraph(f"挂起中：{sessions.get('suspended_sessions', 0)}")
            doc.add_paragraph()
        
        # CPU 和内存
        if 'cpu_memory' in self.metrics:
            doc.add_heading('4.8 CPU 和内存使用情况', level=2)
            mem = self.metrics['cpu_memory']
            total_mem = mem.get('total_memory_kb', 0) / 1024
            target_mem = mem.get('target_memory_kb', 0) / 1024
            ple = mem.get('page_life_expectancy', 0)
            
            doc.add_paragraph(f"总内存使用：{total_mem:.2f} MB")
            doc.add_paragraph(f"目标内存：{target_mem:.2f} MB")
            doc.add_paragraph(f"页面预期寿命 (PLE): {ple} 秒")
            doc.add_paragraph()
        
        # 备份状态
        if 'backup_status' in self.metrics and self.metrics['backup_status']:
            doc.add_heading('4.9 最近备份记录', level=2)
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '数据库'
            hdr_cells[1].text = '类型'
            hdr_cells[2].text = '开始时间'
            hdr_cells[3].text = '完成时间'
            hdr_cells[4].text = '耗时 (秒)'
            hdr_cells[5].text = '大小 (MB)'
            
            for cell in hdr_cells:
                cell.paragraphs[0].runs[0].bold = True
            
            for row in self.metrics['backup_status']:
                row_cells = table.add_row().cells
                row_cells[0].text = str(row.get('database_name', ''))
                row_cells[1].text = str(row.get('backup_type', ''))
                row_cells[2].text = str(row.get('backup_start_date', ''))
                row_cells[3].text = str(row.get('backup_finish_date', ''))
                row_cells[4].text = str(row.get('duration_seconds', 0))
                row_cells[5].text = f"{row.get('backup_size_mb', 0):.2f}"
            doc.add_paragraph()
        
        # 数据库文件
        if 'database_files' in self.metrics and self.metrics['database_files']:
            doc.add_heading('4.10 数据库文件状态', level=2)
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '文件名'
            hdr_cells[1].text = '类型'
            hdr_cells[2].text = '大小 (MB)'
            hdr_cells[3].text = '最大大小 (MB)'
            hdr_cells[4].text = '增长方式'
            hdr_cells[5].text = '状态'
            
            for cell in hdr_cells:
                cell.paragraphs[0].runs[0].bold = True
            
            for row in self.metrics['database_files']:
                row_cells = table.add_row().cells
                row_cells[0].text = str(row.get('file_name', ''))
                row_cells[1].text = str(row.get('file_type', ''))
                row_cells[2].text = f"{row.get('size_mb', 0):.2f}"
                row_cells[3].text = f"{row.get('max_size_mb', 0):.2f}" if row.get('max_size_mb', 0) > 0 else '无限制'
                growth = f"{row.get('growth', 0)}%" if row.get('is_percent_growth', 0) else f"{row.get('growth', 0) * 8 / 1024:.2f} MB"
                row_cells[4].text = growth
                row_cells[5].text = str(row.get('state_desc', ''))
            doc.add_paragraph()
        
        # 锁和死锁
        doc.add_heading('4.11 锁和死锁统计', level=2)
        if 'locks' in self.metrics and self.metrics['locks']:
            for row in self.metrics['locks']:
                doc.add_paragraph(f"数据库 {row.get('database_name', '')}: {row.get('lock_count', 0)} 个锁等待")
        
        if 'deadlocks' in self.metrics:
            deadlock_count = self.metrics['deadlocks'].get('deadlock_count', 0)
            doc.add_paragraph(f"死锁数：{deadlock_count} 次/秒")
        doc.add_paragraph()
        
        # 保存文档
        doc.save(output_path)
        print(f"✓ 报告已生成：{output_path}")
        return output_path


def main():
    """主函数"""
    print("=" * 60)
    print("SQL Server 数据库健康检查工具")
    print("=" * 60)
    
    # 直接配置数据库连接参数，无需手动输入
    server = "10.0.8.162,6037"
    database = "master"
    username = "sa"
    password = "Wg@123456.."
    
    # 创建检查器实例
    checker = SQLServerHealthChecker(server, database, username, password)
    
    # 连接数据库
    if not checker.connect():
        return
    
    try:
        # 执行所有检查
        checker.run_all_checks()
        
        # 生成报告（使用默认文件名）
        report_path = checker.generate_report()
        
        # 显示结果
        print("\n" + "=" * 60)
        print("检查完成！")
        print(f"健康评分：{checker.get_health_score()} 分")
        level, _ = checker.get_health_level()
        print(f"健康等级：{level}")
        print(f"发现问题：{len(checker.health_issues)} 个")
        print(f"报告文件：{report_path}")
        print("=" * 60)
        
    finally:
        checker.disconnect()


if __name__ == "__main__":
    main()
